"""Tests for the Understand My Policy feature: form validation, view, and text extraction."""

import io
import json

from django.core.files.uploadedfile import SimpleUploadedFile
from django.test import TestCase, override_settings
from django.urls import reverse

from fighthealthinsurance.chat_forms import (
    SUPPORTED_POLICY_CONTENT_TYPES,
    SUPPORTED_POLICY_EXTENSIONS,
    UnderstandPolicyForm,
)
from fighthealthinsurance.chat_interface import _detect_policy_analysis_request
from fighthealthinsurance.ml.ml_policy_doc_helper import MLPolicyDocHelper
from fighthealthinsurance.models import PolicyDocument, PolicyDocumentAnalysis


# --- Form validation tests ---


class UnderstandPolicyFormValidationTest(TestCase):
    """Test UnderstandPolicyForm validation for multi-format file support."""

    def _base_data(self, **overrides):
        data = {
            "document_type": "summary_of_benefits",
            "first_name": "Test",
            "last_name": "User",
            "email": "test@example.com",
            "tos_agreement": True,
            "privacy_policy": True,
        }
        data.update(overrides)
        return data

    def _make_pdf(self, name="policy.pdf", size=1024):
        content = b"%PDF-1.4 " + b"x" * (size - 9)
        return SimpleUploadedFile(name, content, content_type="application/pdf")

    def _make_docx(self, name="policy.docx", size=1024):
        # DOCX files are ZIP-based; start with PK magic bytes
        content = b"PK\x03\x04" + b"x" * (size - 4)
        return SimpleUploadedFile(
            name,
            content,
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

    def _make_txt(self, name="policy.txt", size=1024):
        content = b"Policy text content " * (size // 20)
        return SimpleUploadedFile(name, content, content_type="text/plain")

    # --- Accepted file types ---

    def test_accepts_pdf(self):
        form = UnderstandPolicyForm(
            data=self._base_data(),
            files={"policy_document": self._make_pdf()},
        )
        self.assertTrue(form.is_valid(), form.errors)

    def test_accepts_docx(self):
        form = UnderstandPolicyForm(
            data=self._base_data(),
            files={"policy_document": self._make_docx()},
        )
        self.assertTrue(form.is_valid(), form.errors)

    def test_accepts_txt(self):
        form = UnderstandPolicyForm(
            data=self._base_data(),
            files={"policy_document": self._make_txt()},
        )
        self.assertTrue(form.is_valid(), form.errors)

    def test_accepts_rtf(self):
        content = b"{\\rtf1 hello}"
        f = SimpleUploadedFile("policy.rtf", content, content_type="application/rtf")
        form = UnderstandPolicyForm(
            data=self._base_data(),
            files={"policy_document": f},
        )
        self.assertTrue(form.is_valid(), form.errors)

    def test_accepts_doc(self):
        content = b"Some doc content"
        f = SimpleUploadedFile("policy.doc", content, content_type="application/msword")
        form = UnderstandPolicyForm(
            data=self._base_data(),
            files={"policy_document": f},
        )
        self.assertTrue(form.is_valid(), form.errors)

    # --- Rejected file types ---

    def test_rejects_exe(self):
        f = SimpleUploadedFile("malware.exe", b"MZ" + b"\x00" * 100, content_type="application/octet-stream")
        form = UnderstandPolicyForm(
            data=self._base_data(),
            files={"policy_document": f},
        )
        self.assertFalse(form.is_valid())
        self.assertIn("policy_document", form.errors)

    def test_rejects_html(self):
        f = SimpleUploadedFile("page.html", b"<html></html>", content_type="text/html")
        form = UnderstandPolicyForm(
            data=self._base_data(),
            files={"policy_document": f},
        )
        self.assertFalse(form.is_valid())
        self.assertIn("policy_document", form.errors)

    def test_rejects_jpg(self):
        f = SimpleUploadedFile("photo.jpg", b"\xff\xd8\xff" + b"\x00" * 100, content_type="image/jpeg")
        form = UnderstandPolicyForm(
            data=self._base_data(),
            files={"policy_document": f},
        )
        self.assertFalse(form.is_valid())
        self.assertIn("policy_document", form.errors)

    # --- Magic byte validation ---

    def test_rejects_fake_pdf_extension(self):
        """A file named .pdf but without PDF magic bytes should be rejected."""
        f = SimpleUploadedFile("fake.pdf", b"not a pdf at all", content_type="application/pdf")
        form = UnderstandPolicyForm(
            data=self._base_data(),
            files={"policy_document": f},
        )
        self.assertFalse(form.is_valid())
        self.assertIn("policy_document", form.errors)

    def test_rejects_fake_docx_extension(self):
        """A file named .docx but without PK magic bytes should be rejected."""
        f = SimpleUploadedFile("fake.docx", b"not a docx", content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        form = UnderstandPolicyForm(
            data=self._base_data(),
            files={"policy_document": f},
        )
        self.assertFalse(form.is_valid())
        self.assertIn("policy_document", form.errors)

    # --- Size validation ---

    def test_rejects_oversized_file(self):
        """Files over 20MB should be rejected."""
        # Create a file just over 20MB
        big_content = b"%PDF-1.4 " + b"x" * (20 * 1024 * 1024 + 1)
        f = SimpleUploadedFile("big.pdf", big_content, content_type="application/pdf")
        form = UnderstandPolicyForm(
            data=self._base_data(),
            files={"policy_document": f},
        )
        self.assertFalse(form.is_valid())
        self.assertIn("policy_document", form.errors)

    # --- Required fields ---

    def test_requires_tos(self):
        form = UnderstandPolicyForm(
            data=self._base_data(tos_agreement=False),
            files={"policy_document": self._make_pdf()},
        )
        self.assertFalse(form.is_valid())
        self.assertIn("tos_agreement", form.errors)

    def test_requires_email(self):
        data = self._base_data()
        data.pop("email")
        form = UnderstandPolicyForm(
            data=data,
            files={"policy_document": self._make_pdf()},
        )
        self.assertFalse(form.is_valid())
        self.assertIn("email", form.errors)

    def test_user_question_optional(self):
        form = UnderstandPolicyForm(
            data=self._base_data(),
            files={"policy_document": self._make_pdf()},
        )
        self.assertTrue(form.is_valid(), form.errors)
        self.assertEqual(form.cleaned_data["user_question"], "")


# --- Text extraction tests ---


class TextExtractionTest(TestCase):
    """Test MLPolicyDocHelper text extraction for multiple file formats."""

    def test_extract_text_from_real_pdf(self):
        """Create a minimal PDF with pymupdf and verify extraction."""
        import pymupdf

        doc = pymupdf.open()
        page = doc.new_page()
        page.insert_text((72, 72), "This is test coverage text for page one.")
        pdf_bytes = doc.tobytes()
        doc.close()

        full_text, page_dict = MLPolicyDocHelper._extract_text_from_pdf_bytes(pdf_bytes)
        self.assertIn("test coverage text", full_text)
        self.assertIn(1, page_dict)
        self.assertIn("test coverage text", page_dict[1])

    def test_extract_text_from_docx(self):
        """Create a minimal DOCX and verify extraction."""
        import docx

        doc = docx.Document()
        doc.add_paragraph("Summary of Benefits and Coverage")
        doc.add_paragraph("This plan covers preventive care at no cost.")
        buffer = io.BytesIO()
        doc.save(buffer)
        docx_bytes = buffer.getvalue()

        full_text, page_dict = MLPolicyDocHelper._extract_text_from_docx_bytes(docx_bytes)
        self.assertIn("Summary of Benefits", full_text)
        self.assertIn("preventive care", full_text)
        self.assertTrue(len(page_dict) >= 1)

    def test_extract_text_from_plaintext(self):
        """Verify extraction from plain text bytes."""
        text_bytes = (
            b"Exclusions: Cosmetic surgery is not covered.\n"
            b"Appeals: You may appeal within 180 days.\n"
        )

        full_text, page_dict = MLPolicyDocHelper._extract_text_from_plaintext_bytes(
            text_bytes
        )
        self.assertIn("Cosmetic surgery", full_text)
        self.assertIn("180 days", full_text)
        self.assertIn(1, page_dict)

    def test_extract_text_dispatches_by_extension(self):
        """The extract_text_from_bytes() dispatcher picks the right method."""
        text_bytes = b"Hello from txt dispatch test"

        full_text, _ = MLPolicyDocHelper.extract_text_from_bytes(text_bytes, "test.txt")
        self.assertIn("Hello from txt dispatch test", full_text)

    def test_extract_text_unsupported_extension(self):
        """Unsupported extensions return empty results."""
        full_text, page_dict = MLPolicyDocHelper.extract_text_from_bytes(
            b"some content", "file.xyz"
        )
        self.assertEqual(full_text, "")
        self.assertEqual(page_dict, {})

    def test_extract_text_corrupt_pdf(self):
        """Corrupt PDF returns empty without crashing."""
        full_text, page_dict = MLPolicyDocHelper._extract_text_from_pdf_bytes(
            b"this is not a pdf"
        )
        self.assertEqual(full_text, "")
        self.assertEqual(page_dict, {})


# --- JSON parsing tests ---


class AnalysisResponseParsingTest(TestCase):
    """Test MLPolicyDocHelper._parse_analysis_response for robustness."""

    def test_parse_valid_json(self):
        response = json.dumps(
            {
                "exclusions": [{"text": "No cosmetic", "page": 1, "explanation": ""}],
                "inclusions": [{"text": "Preventive care", "page": 2, "explanation": ""}],
                "appeal_clauses": [{"text": "180 days", "page": 5, "type": "appeal_deadline"}],
                "summary": "Basic coverage plan.",
                "quotable_sections": [],
            }
        )
        result = MLPolicyDocHelper._parse_analysis_response(response)
        self.assertIsNotNone(result)
        self.assertEqual(len(result["exclusions"]), 1)
        self.assertEqual(result["summary"], "Basic coverage plan.")

    def test_parse_json_embedded_in_text(self):
        response = 'Here is the analysis:\n```json\n{"exclusions": [], "inclusions": [], "appeal_clauses": [], "summary": "Test", "quotable_sections": []}\n```'
        result = MLPolicyDocHelper._parse_analysis_response(response)
        self.assertIsNotNone(result)
        self.assertEqual(result["summary"], "Test")

    def test_parse_missing_required_keys(self):
        response = json.dumps({"exclusions": [], "summary": "incomplete"})
        result = MLPolicyDocHelper._parse_analysis_response(response)
        self.assertIsNone(result)

    def test_parse_garbage_input(self):
        result = MLPolicyDocHelper._parse_analysis_response("not json at all")
        self.assertIsNone(result)

    def test_parse_empty_string(self):
        result = MLPolicyDocHelper._parse_analysis_response("")
        self.assertIsNone(result)

    def test_parse_json_with_trailing_text(self):
        """raw_decode should handle JSON followed by extra text."""
        valid_json = json.dumps(
            {
                "exclusions": [],
                "inclusions": [],
                "appeal_clauses": [],
                "summary": "Good",
                "quotable_sections": [],
            }
        )
        response = f"Analysis: {valid_json}\n\nHope this helps! }}"
        result = MLPolicyDocHelper._parse_analysis_response(response)
        self.assertIsNotNone(result)
        self.assertEqual(result["summary"], "Good")


# --- Policy analysis detection tests ---


class PolicyAnalysisDetectionTest(TestCase):
    """Test the chat interface regex for detecting policy analysis requests."""

    def test_matches_structured_form_message(self):
        """The structured message from UnderstandPolicyView should match."""
        msg = (
            "I've uploaded my insurance Summary of Benefits: my_policy.pdf\n\n"
            "My question: Is my MRI covered?\n\n"
            "Please analyze this document and help me understand:\n"
            "1. What is covered (inclusions)\n"
            "2. What is NOT covered (exclusions)\n"
            "3. Any clauses that would be useful if I need to appeal a denial\n"
            "4. Key quotes I could use in an appeal letter"
        )
        self.assertTrue(_detect_policy_analysis_request(msg))

    def test_matches_without_question(self):
        """Form message without optional question should still match."""
        msg = (
            "I've uploaded my insurance Medical Policy: plan.docx\n\n"
            "Please analyze this document and help me understand:\n"
            "1. What is covered"
        )
        self.assertTrue(_detect_policy_analysis_request(msg))

    def test_no_false_positive_on_normal_chat(self):
        """Normal chat messages should NOT trigger policy analysis."""
        normal_messages = [
            "What is covered under my plan?",
            "Can you explain the summary of benefits?",
            "My medical policy says cosmetic surgery is excluded.",
            "What is not covered?",
            "Tell me about insurance documents.",
            "I need help with my denied claim for an MRI.",
            "The exclusions section mentions experimental treatments.",
        ]
        for msg in normal_messages:
            self.assertFalse(
                _detect_policy_analysis_request(msg),
                f"False positive on: {msg!r}",
            )


# --- View tests ---


@override_settings(RECAPTCHA_TESTING=True)
class UnderstandPolicyViewTest(TestCase):
    """Test the UnderstandPolicyView GET and POST."""

    def test_get_returns_form(self):
        response = self.client.get(reverse("understand_policy"))
        self.assertEqual(response.status_code, 200)
        self.assertContains(response, "Understand My Policy")
        self.assertContains(response, "policy_document")

    def test_post_invalid_no_file(self):
        response = self.client.post(
            reverse("understand_policy"),
            data={
                "document_type": "summary_of_benefits",
                "first_name": "Test",
                "last_name": "User",
                "email": "test@example.com",
                "tos_agreement": True,
                "privacy_policy": True,
            },
        )
        self.assertEqual(response.status_code, 200)  # Re-renders form
        self.assertContains(response, "This field is required")

    def test_post_valid_pdf_creates_document(self):
        pdf_content = b"%PDF-1.4 test content"
        f = SimpleUploadedFile("test_policy.pdf", pdf_content, content_type="application/pdf")
        response = self.client.post(
            reverse("understand_policy"),
            data={
                "document_type": "summary_of_benefits",
                "user_question": "Is my MRI covered?",
                "first_name": "Test",
                "last_name": "User",
                "email": "test@example.com",
                "tos_agreement": True,
                "privacy_policy": True,
                "policy_document": f,
            },
        )
        # Should render chat_redirect template (200, not 302)
        self.assertEqual(response.status_code, 200)
        # A PolicyDocument should have been created
        self.assertEqual(PolicyDocument.objects.count(), 1)
        doc = PolicyDocument.objects.first()
        self.assertEqual(doc.document_type, "summary_of_benefits")
        self.assertIn("test_policy", doc.filename)

    def test_post_valid_docx_creates_document(self):
        docx_content = b"PK\x03\x04" + b"x" * 100
        f = SimpleUploadedFile(
            "test_policy.docx",
            docx_content,
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
        response = self.client.post(
            reverse("understand_policy"),
            data={
                "document_type": "medical_policy",
                "first_name": "Test",
                "last_name": "User",
                "email": "test@example.com",
                "tos_agreement": True,
                "privacy_policy": True,
                "policy_document": f,
            },
        )
        self.assertEqual(response.status_code, 200)
        self.assertEqual(PolicyDocument.objects.count(), 1)
        doc = PolicyDocument.objects.first()
        self.assertEqual(doc.document_type, "medical_policy")

    def test_post_stores_document_id_in_session(self):
        pdf_content = b"%PDF-1.4 test"
        f = SimpleUploadedFile("policy.pdf", pdf_content, content_type="application/pdf")
        self.client.post(
            reverse("understand_policy"),
            data={
                "document_type": "other",
                "first_name": "Test",
                "last_name": "User",
                "email": "test@example.com",
                "tos_agreement": True,
                "privacy_policy": True,
                "policy_document": f,
            },
        )
        session = self.client.session
        self.assertIn("policy_document_id", session)
        # The stored ID should match the created document
        doc = PolicyDocument.objects.first()
        self.assertEqual(session["policy_document_id"], str(doc.id))


# --- Model tests ---


class PolicyDocumentModelTest(TestCase):
    """Test PolicyDocument and PolicyDocumentAnalysis models."""

    def test_create_policy_document(self):
        doc = PolicyDocument.objects.create(
            document_type="summary_of_benefits",
            filename="test.pdf",
            hashed_email="abc123",
            session_key="sess123",
        )
        self.assertIsNotNone(doc.id)
        self.assertEqual(str(doc), "PolicyDocument: test.pdf (summary_of_benefits)")

    def test_create_analysis(self):
        doc = PolicyDocument.objects.create(
            document_type="other",
            filename="plan.docx",
        )
        analysis = PolicyDocumentAnalysis.objects.create(
            policy_document=doc,
            user_question="Is my MRI covered?",
            exclusions=[{"text": "No cosmetic", "page": 1}],
            inclusions=[{"text": "Preventive care", "page": 2}],
            appeal_clauses=[],
            summary="Basic plan.",
        )
        self.assertIsNotNone(analysis.id)
        self.assertEqual(analysis.policy_document, doc)

    def test_cascade_delete(self):
        """Deleting a PolicyDocument should cascade-delete its analyses."""
        doc = PolicyDocument.objects.create(filename="test.pdf")
        PolicyDocumentAnalysis.objects.create(
            policy_document=doc,
            summary="Test",
        )
        self.assertEqual(PolicyDocumentAnalysis.objects.count(), 1)
        doc.delete()
        self.assertEqual(PolicyDocumentAnalysis.objects.count(), 0)


# --- Format analysis for chat tests ---


class FormatAnalysisForChatTest(TestCase):
    """Test MLPolicyDocHelper.format_analysis_for_chat output."""

    def test_format_with_all_sections(self):
        doc = PolicyDocument.objects.create(filename="test.pdf")
        analysis = PolicyDocumentAnalysis.objects.create(
            policy_document=doc,
            summary="This plan provides basic coverage.",
            exclusions=[{"text": "Cosmetic surgery", "page": 3, "explanation": "Not covered"}],
            inclusions=[{"text": "Emergency room visits", "page": 1, "explanation": "Covered"}],
            appeal_clauses=[{"text": "Appeal within 180 days", "page": 10, "type": "appeal_deadline"}],
            quotable_sections=[{"quote": "All emergency services are covered", "page": 1, "relevance": "ER coverage"}],
        )
        output = MLPolicyDocHelper.format_analysis_for_chat(analysis)
        self.assertIn("## Summary", output)
        self.assertIn("basic coverage", output)
        self.assertIn("Cosmetic surgery", output)
        self.assertIn("Emergency room", output)
        self.assertIn("180 days", output)
        self.assertIn("Disclaimer", output)

    def test_format_without_disclaimer(self):
        doc = PolicyDocument.objects.create(filename="test.pdf")
        analysis = PolicyDocumentAnalysis.objects.create(
            policy_document=doc,
            summary="Short summary.",
        )
        output = MLPolicyDocHelper.format_analysis_for_chat(
            analysis, include_disclaimer=False
        )
        self.assertNotIn("Disclaimer", output)
